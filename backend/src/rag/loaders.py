"""Multi-format document loaders for KnowBase.

Supports .txt, .md, .pdf, .docx, and .html files.
Each loader returns LangChain Documents with ``source`` (and optionally ``page``)
metadata pre-filled. Downstream splitting (``_prepare_splits``) appends the
standard chunk-level metadata.
"""

from __future__ import annotations

import ipaddress
import socket
from pathlib import Path
from typing import List, Optional
from urllib.parse import urljoin, urlparse

import requests
from langchain_core.documents import Document
from requests.adapters import HTTPAdapter
from urllib3 import HTTPConnectionPool, HTTPSConnectionPool
from urllib3.connection import HTTPConnection, HTTPSConnection
from urllib3.poolmanager import PoolManager


from src.rag.models import normalize_source


def load_document(file_path: str, source_name: str | None = None) -> List[Document]:
    """Load a document by file extension and return LangChain Documents."""
    path = Path(file_path)
    ext = path.suffix.lower()
    display_source = normalize_source(source_name or path.name)

    loader = _LOADER_MAP.get(ext)
    if loader is None:
        raise ValueError(f"不支持的文件格式：{ext}")

    docs = loader(str(path))
    for doc in docs:
        doc.metadata["source"] = display_source
    return docs


def _load_text(file_path: str) -> List[Document]:
    with open(file_path, encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": Path(file_path).name})]


def _load_pdf(file_path: str) -> List[Document]:
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    docs: List[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            docs.append(
                Document(
                    page_content=text.strip(),
                    metadata={"source": Path(file_path).name, "page": i + 1},
                )
            )
    if not docs:
        docs.append(
            Document(page_content="", metadata={"source": Path(file_path).name, "page": 1})
        )
    return docs


def _load_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs) if paragraphs else ""
    return [Document(page_content=text, metadata={"source": Path(file_path).name})]


def _load_html(file_path: str) -> List[Document]:
    from bs4 import BeautifulSoup
    with open(file_path, encoding="utf-8") as f:
        content = f.read()
    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    try:
        from markdownify import markdownify as md
        text = md(str(soup), heading_style="ATX")
    except ImportError:
        text = soup.get_text(separator="\n", strip=True)
    return [Document(page_content=text.strip(), metadata={"source": Path(file_path).name})]


_LOADER_MAP = {
    ".txt": _load_text,
    ".md": _load_text,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".html": _load_html,
    ".htm": _load_html,
}


# ---------------------------------------------------------------------------
# SSRF 防护 —— 自定义 HTTPConnection / HTTPSConnection + PoolManager
# ---------------------------------------------------------------------------
# 重写 _new_conn() 在 TCP 连接层使用验证过的 IP。
# host/_dns_host 保持原始域名不变，TLS SNI、证书校验正确。
# 不修改全局状态，不依赖 monkey-patch，完全并发安全。
# ---------------------------------------------------------------------------

def _is_safe_ip(addr: ipaddress.IPv4Address | ipaddress.IPv6Address) -> bool:
    """IP 是否安全（必须 is_global 且不是组播地址）。"""
    return addr.is_global and not addr.is_multicast


def _validate_url(url: str) -> Optional[str]:
    """解析域名，验证所有解析 IP 安全，返回第一个公网 IP。

    Returns:
        str — 安全的公网 IP
        None — IP 直连（已过检），无需 pin
    """
    parsed = urlparse(url)
    host = parsed.hostname
    if not host:
        return None

    # IP 直连
    try:
        addr = ipaddress.ip_address(host)
    except ValueError:
        pass
    else:
        if not _is_safe_ip(addr):
            raise ValueError("不允许访问内网地址（SSRF 防护）")
        return None

    # 域名 → 解析并检查所有结果
    seen: set[str] = set()
    resolved: list[str] = []
    try:
        for _, _, _, _, sockaddr in socket.getaddrinfo(host, None):
            ip = sockaddr[0]
            if ip not in seen:
                seen.add(ip)
                resolved.append(ip)
    except OSError:
        raise ValueError("DNS 解析失败")

    if not resolved:
        raise ValueError("DNS 解析失败：无可用地址")

    for ip in resolved:
        addr = ipaddress.ip_address(ip)
        if not _is_safe_ip(addr):
            raise ValueError("不允许访问内网地址（SSRF 防护）")

    return resolved[0]


def _is_private_url(url: str) -> bool:
    """（已废弃，保留仅为兼容已有测试引用）"""
    try:
        _validate_url(url)
        return False
    except ValueError:
        return True


# 共用 Connection 构造逻辑

def _make_pinned_connection_cls(base_cls, pinned_ip: str):
    """返回一个 Connection 子类，重写 _new_conn() 在 TCP 层用 pinned_ip 连接。

    不修改 self.host / self._dns_host，保证 TLS SNI、证书校验、
    Host header 都使用原始域名。仅 TCP connect 的目标地址改为 pinned_ip。
    """
    class _PinnedConnection(base_cls):  # type: ignore[valid-type]
        def __init__(self, host, port=None, **kwargs):
            self._pinned_ip = pinned_ip
            super().__init__(host, port, **kwargs)

        def _new_conn(self):
            from urllib3.util.connection import create_connection

            return create_connection(
                (self._pinned_ip, self.port),
                timeout=self.timeout,
                source_address=self.source_address,
                socket_options=self.socket_options,
            )

    _PinnedConnection.__name__ = f"_Pinned{base_cls.__name__}"
    return _PinnedConnection


def _make_pinned_pool_manager(pin_ip: str, connections, maxsize, block, **pool_kwargs):
    """返回一个 PoolManager，HTTP 和 HTTPS 池都使用 pin 后的 Connection。"""
    https_cls = _make_pinned_connection_cls(HTTPSConnection, pin_ip)
    http_cls = _make_pinned_connection_cls(HTTPConnection, pin_ip)

    class _PinnedHTTPSConnectionPool(HTTPSConnectionPool):
        ConnectionCls = https_cls

    class _PinnedHTTPConnectionPool(HTTPConnectionPool):
        ConnectionCls = http_cls

    class _PinnedPool(PoolManager):
        def _new_pool(self, scheme, host, port, request_context=None):
            if scheme == "https":
                return _PinnedHTTPSConnectionPool(host, port, **self.connection_pool_kw)
            return _PinnedHTTPConnectionPool(host, port, **self.connection_pool_kw)

    return _PinnedPool(connections, maxsize=maxsize, block=block, **pool_kwargs)


class _PinnedIPAdapter(HTTPAdapter):
    """HTTPAdapter：HTTP 和 HTTPS 连接都 pin 到指定 IP。"""

    def __init__(self, pinned_ip, *args, **kwargs):
        self._pinned_ip = pinned_ip
        super().__init__(*args, **kwargs)

    def init_poolmanager(self, connections, maxsize, block=False, **pool_kwargs):
        self.poolmanager = _make_pinned_pool_manager(
            self._pinned_ip, connections, maxsize, block, **pool_kwargs
        )


def _make_pinned_request(url: str, headers: dict, pin_ip: str) -> requests.Response:
    """用自定义 Adapter 将请求 pin 到指定 IP。

    HTTP/HTTPS 都通过自定义 Connection 重写 TCP connect 目标实现，
    不依赖全局状态，完全并发安全。
    """
    with requests.Session() as session:
        adapter = _PinnedIPAdapter(pinned_ip=pin_ip)
        session.mount("https://", adapter)
        session.mount("http://", adapter)
        return session.get(url, headers=headers, timeout=30, allow_redirects=False)


# ---------------------------------------------------------------------------
# Public URL loader
# ---------------------------------------------------------------------------

def load_url(url: str, source_name: str | None = None) -> List[Document]:
    """Fetch a URL and extract its main content as LangChain Documents.

    SSRF 防护：先 DNS 验证所有 IP 安全，再用自定义 TransportAdapter
    在连接层面 pin 到已验证 IP（防 DNS rebinding）。
    URL 保持原始域名 → TLS SNI、Host header、证书校验全部正确。
    """
    from bs4 import BeautifulSoup

    display_source = source_name or url
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )
    }

    def _do_request(target_url: str) -> requests.Response:
        pin_ip = _validate_url(target_url)
        if pin_ip is not None:
            return _make_pinned_request(target_url, headers, pin_ip)
        return requests.get(target_url, headers=headers, timeout=30, allow_redirects=False)

    # 逐跳跟随重定向，每跳都做 IP pin
    target = url
    max_redirects = 10
    for _ in range(max_redirects + 1):
        resp = _do_request(target)
        if 300 <= resp.status_code < 400:
            location = resp.headers.get("Location")
            if not location:
                break
            target = urljoin(target, location)
            continue
        break
    else:
        raise ValueError("重定向次数过多")

    resp.raise_for_status()
    resp.encoding = resp.apparent_encoding or "utf-8"

    if "accounts.feishu.cn" in target or "/login" in target:
        raise ValueError("目标页面跳转到了登录页。请确认该 URL 可公开访问，或使用可直接访问的网页链接。")

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
        tag.decompose()

    main = (
        soup.find("article")
        or soup.find("main")
        or soup.find(class_=lambda c: c and any(k in (c or "").lower() for k in ("content", "article", "post")))
        or soup.body
    )

    try:
        from markdownify import markdownify as md
        text = md(str(main or soup), heading_style="ATX")
    except ImportError:
        text = (main or soup).get_text(separator="\n", strip=True)

    text = text.strip()
    if not text:
        raise ValueError("网页抓取成功，但未提取到正文内容")

    title_tag = soup.find("title")
    title = title_tag.get_text(strip=True) if title_tag else url

    return [
        Document(
            page_content=text,
            metadata={"source": display_source, "title": title, "url": url},
        )
    ]
