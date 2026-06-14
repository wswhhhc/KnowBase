"""Multi-format document loaders for KnowBase.

Supports .txt, .md, .pdf, .docx, and .html files.
Each loader returns LangChain Documents with ``source`` (and optionally ``page``)
metadata pre-filled. Downstream splitting (``_prepare_splits``) appends the
standard chunk-level metadata.
"""

from __future__ import annotations

from pathlib import Path
from typing import List

from langchain_core.documents import Document


def load_document(file_path: str, source_name: str | None = None) -> List[Document]:
    """Load a document by file extension and return LangChain Documents.

    Parameters
    ----------
    file_path:
        Absolute or relative path to the file.
    source_name:
        Optional display name; defaults to the basename of ``file_path``.
        This becomes each Document's ``source`` metadata.

    Returns
    -------
    List[Document]
        One or more Documents. For PDFs there is one Document per page (with
        ``page`` metadata). Other formats produce a single Document.
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    display_source = Path(source_name or path.name).name

    loader = _LOADER_MAP.get(ext)
    if loader is None:
        raise ValueError(f"不支持的文件格式：{ext}")

    docs = loader(str(path))
    for doc in docs:
        doc.metadata["source"] = display_source
    return docs


# ---------------------------------------------------------------------------
# Internal loaders
# ---------------------------------------------------------------------------

def _load_text(file_path: str) -> List[Document]:
    from langchain_community.document_loaders import TextLoader

    loader = TextLoader(file_path, encoding="utf-8")
    return loader.load()


def _load_pdf(file_path: str) -> List[Document]:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    docs: List[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            docs.append(
                Document(
                    page_content=text.strip(),
                    metadata={"source": Path(file_path).name, "page": i + 1},
                )
            )
    if not docs:
        # Return an empty document so the caller sees a known format.
        docs.append(
            Document(
                page_content="",
                metadata={"source": Path(file_path).name, "page": 1},
            )
        )
    return docs


def _load_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs) if paragraphs else ""
    return [Document(page_content=text, metadata={"source": Path(file_path).name})]


def _load_html(file_path: str) -> List[Document]:
    from bs4 import BeautifulSoup

    with open(file_path, encoding="utf-8") as f:
        content = f.read()

    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    try:
        from markdownify import markdownify as md

        text = md(str(soup), heading_style="ATX")
    except ImportError:
        text = soup.get_text(separator="\n", strip=True)

    return [Document(page_content=text.strip(), metadata={"source": Path(file_path).name})]


_LOADER_MAP = {
    ".txt": _load_text,
    ".md": _load_text,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".html": _load_html,
    ".htm": _load_html,
}


def load_url(url: str, source_name: str | None = None) -> List[Document]:
    """Fetch a URL and extract its main content as LangChain Documents.

    Strips navigation, ads, and other boilerplate via BeautifulSoup.
    The URL itself is stored as the ``source`` metadata.
    """
    import requests
    from bs4 import BeautifulSoup

    display_source = source_name or url
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )
    }
    resp = requests.get(url, headers=headers, timeout=30)
    resp.raise_for_status()
    resp.encoding = resp.apparent_encoding or "utf-8"

    final_url = getattr(resp, "url", url) or url
    if "accounts.feishu.cn" in final_url or "/login" in final_url:
        raise ValueError("目标页面跳转到了登录页。请确认该 URL 可公开访问，或使用可直接访问的网页链接。")

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
        tag.decompose()

    # Try to find main content area
    main = (
        soup.find("article")
        or soup.find("main")
        or soup.find(class_=lambda c: c and any(k in (c or "").lower() for k in ("content", "article", "post")))
        or soup.body
    )

    try:
        from markdownify import markdownify as md

        text = md(str(main or soup), heading_style="ATX")
    except ImportError:
        text = (main or soup).get_text(separator="\n", strip=True)

    text = text.strip()
    if not text:
        raise ValueError("网页抓取成功，但未提取到正文内容。该页面可能依赖登录、前端渲染，或正文不可直接抓取。")

    title_tag = soup.find("title")
    title = title_tag.get_text(strip=True) if title_tag else url

    return [
        Document(
            page_content=text,
            metadata={"source": display_source, "title": title, "url": url},
        )
    ]
